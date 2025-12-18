"""
Document Indexer - Handles chunking, embedding, and storage in ChromaDB.

Includes anchor string injection for BM25 boost on specific references like:
- "Figure 37" → [anchor: figure 37]
- "Table 2.3" → [anchor: table 2.3]
- "Section 4.1" → [anchor: section 4.1]
"""
import logging
import re
from typing import List, Optional, Set
from datetime import datetime
import uuid

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter  # type: ignore[import-untyped]
from langchain_openai import OpenAIEmbeddings
from langchain_chroma import Chroma  # type: ignore[import-untyped]

from .config import RAGConfig, get_config
from .loaders import DocumentLoaderFactory, LoaderResult

logger = logging.getLogger(__name__)


class DocumentIndexer:
    """
    Handles document ingestion pipeline:
    1. Load documents from various sources
    2. Split into chunks
    3. Generate embeddings
    4. Store in ChromaDB
    """
    
    def __init__(self, config: Optional[RAGConfig] = None):
        """Initialize the indexer with configuration."""
        self.config = config or get_config()
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.config.chunk_size,
            chunk_overlap=self.config.chunk_overlap,
            add_start_index=True,  # Track position in original document
            separators=["\n\n", "\n", ". ", " ", ""],
        )
        
        # Initialize embeddings model
        self.embeddings = OpenAIEmbeddings(
            model=self.config.embedding_model,
            openai_api_key=self.config.openai_api_key,
        )
        
        # Initialize or load vector store
        self.vector_store = Chroma(
            collection_name=self.config.collection_name,
            embedding_function=self.embeddings,
            persist_directory=self.config.persist_directory,
        )
        
        logger.info(f"DocumentIndexer initialized with persist_directory: {self.config.persist_directory}")
    
    def ingest(
        self,
        source: str,
        title: Optional[str] = None,
        document_id: Optional[str] = None,
    ) -> dict:
        """
        Ingest a document from a source (URL or file path).
        
        Args:
            source: URL or file path to ingest
            title: Optional title for the document
            document_id: Optional unique ID for the document
            
        Returns:
            Dictionary with ingestion results
        """
        try:
            # Generate document ID if not provided
            if document_id is None:
                document_id = str(uuid.uuid4())[:8]
            
            logger.info(f"Starting ingestion for: {source} (ID: {document_id})")
            
            # Load documents
            loader_result = DocumentLoaderFactory.load(source, title=title)
            
            # Split into chunks
            chunks = self._split_documents(loader_result.documents, document_id)
            
            # Add to vector store
            ids = self._store_chunks(chunks, document_id)
            
            result = {
                "success": True,
                "document_id": document_id,
                "title": loader_result.title,
                "source_type": loader_result.source_type,
                "source_url": loader_result.source_url,
                "num_chunks": len(chunks),
                "chunk_ids": ids,
                "ingested_at": datetime.now().isoformat(),
            }
            
            logger.info(f"Successfully ingested {len(chunks)} chunks from: {source}")
            return result
            
        except Exception as e:
            logger.error(f"Error during ingestion: {e}")
            return {
                "success": False,
                "error": str(e),
                "source": source,
            }
    
    def ingest_text(
        self,
        text: str,
        title: str,
        document_id: Optional[str] = None,
        metadata: Optional[dict] = None,
    ) -> dict:
        """
        Ingest raw text content directly.
        
        Args:
            text: Raw text content to ingest
            title: Title for the document
            document_id: Optional unique ID
            metadata: Optional additional metadata
            
        Returns:
            Dictionary with ingestion results
        """
        try:
            if document_id is None:
                document_id = str(uuid.uuid4())[:8]
            
            # Create document
            doc = Document(
                page_content=text,
                metadata={
                    "source_type": "text",
                    "source_url": f"text://{document_id}",
                    "title": title,
                    "ingested_at": datetime.now().isoformat(),
                    **(metadata or {}),
                }
            )
            
            # Split into chunks
            chunks = self._split_documents([doc], document_id)
            
            # Store chunks
            ids = self._store_chunks(chunks, document_id)
            
            return {
                "success": True,
                "document_id": document_id,
                "title": title,
                "source_type": "text",
                "num_chunks": len(chunks),
                "chunk_ids": ids,
                "ingested_at": datetime.now().isoformat(),
            }
            
        except Exception as e:
            logger.error(f"Error ingesting text: {e}")
            return {
                "success": False,
                "error": str(e),
            }
    
    @staticmethod
    def _extract_anchors(text: str) -> Set[str]:
        """
        Extract structural anchors from text for BM25 boost.
        
        Detects patterns like:
        - Figure 37, Fig. 37, fig 37
        - Table 2.3, Tab. 2
        - Section 4.1, Sec. 4
        - Chapter 7, Ch. 7
        - Equation 5, Eq. 5
        - Reference [15], [15]
        - Page 42
        
        Returns:
            Set of normalized anchor strings like "figure 37", "table 2.3"
        """
        anchors: Set[str] = set()
        text_lower = text.lower()
        
        # Figure patterns
        for match in re.finditer(r"\b(?:figure|fig\.?)\s*(\d+)\b", text_lower):
            anchors.add(f"figure {match.group(1)}")
        
        # Table patterns
        for match in re.finditer(r"\b(?:table|tab\.?)\s*([\d.]+)\b", text_lower):
            anchors.add(f"table {match.group(1)}")
        
        # Section patterns
        for match in re.finditer(r"\b(?:section|sec\.?)\s*([\d.]+)\b", text_lower):
            anchors.add(f"section {match.group(1)}")
        
        # Chapter patterns
        for match in re.finditer(r"\b(?:chapter|ch\.?)\s*(\d+)\b", text_lower):
            anchors.add(f"chapter {match.group(1)}")
        
        # Equation patterns
        for match in re.finditer(r"\b(?:equation|eq\.?)\s*(\d+)\b", text_lower):
            anchors.add(f"equation {match.group(1)}")
        
        # Reference patterns [X]
        for match in re.finditer(r"\[(\d+)\]", text):
            anchors.add(f"reference {match.group(1)}")
        
        # Page patterns (from PDF metadata, typically)
        for match in re.finditer(r"\bpage\s+(\d+)\b", text_lower):
            anchors.add(f"page {match.group(1)}")
        
        return anchors
    
    def _split_documents(
        self,
        documents: List[Document],
        document_id: str,
    ) -> List[Document]:
        """
        Split documents into chunks with proper metadata and anchor strings.
        
        Anchor strings are appended to chunks to boost BM25 keyword matching
        for specific references like "Figure 37", "Table 2.3", etc.
        """
        chunks = self.text_splitter.split_documents(documents)
        
        # Add document_id, chunk_index, and anchor strings to each chunk
        for i, chunk in enumerate(chunks):
            chunk.metadata["document_id"] = document_id
            chunk.metadata["chunk_index"] = i
            chunk.metadata["total_chunks"] = len(chunks)
            
            # Extract and store anchors in metadata
            anchors = self._extract_anchors(chunk.page_content)
            if anchors:
                # Store as comma-separated string (Chroma doesn't accept lists)
                chunk.metadata["anchors"] = ", ".join(sorted(anchors))
                
                # Append anchor strings to content for BM25 boost
                # Format: [anchor: figure 37] [anchor: table 2.3]
                anchor_suffix = " ".join(f"[anchor: {a}]" for a in sorted(anchors))
                chunk.page_content = f"{chunk.page_content}\n\n{anchor_suffix}"
                
                logger.debug(f"Chunk {i}: Added anchors {anchors}")
        
        avg_chars = sum(len(c.page_content) for c in chunks) // len(chunks) if chunks else 0
        anchor_count = sum(1 for c in chunks if c.metadata.get("anchors"))
        logger.info(
            f"Split into {len(chunks)} chunks (avg {avg_chars} chars, "
            f"{anchor_count} with anchors)"
        )
        return chunks
    
    def _store_chunks(
        self,
        chunks: List[Document],
        document_id: str,
    ) -> List[str]:
        """Store chunks in the vector store."""
        # Generate unique IDs for each chunk
        ids = [f"{document_id}_{i}" for i in range(len(chunks))]
        
        # Add to vector store
        self.vector_store.add_documents(chunks, ids=ids)
        
        logger.info(f"Stored {len(chunks)} chunks in vector store")
        return ids
    
    def delete_document(self, document_id: str) -> bool:
        """Delete all chunks for a document from the vector store."""
        try:
            # Get all chunk IDs for this document
            # Note: Chroma doesn't have a direct "delete by metadata" in langchain wrapper
            # We'll need to delete by ID pattern
            collection = self.vector_store._collection
            
            # Get documents with matching document_id
            results = collection.get(
                where={"document_id": document_id}
            )
            
            if results and results["ids"]:
                collection.delete(ids=results["ids"])
                logger.info(f"Deleted {len(results['ids'])} chunks for document: {document_id}")
                return True
            else:
                logger.warning(f"No chunks found for document: {document_id}")
                return False
                
        except Exception as e:
            logger.error(f"Error deleting document: {e}")
            return False
    
    def list_documents(self) -> List[dict]:
        """List all unique documents in the vector store."""
        try:
            collection = self.vector_store._collection
            
            # Get all documents
            results = collection.get(include=["metadatas"])
            
            # Extract unique documents
            documents = {}
            for metadata in results.get("metadatas", []):
                if metadata:
                    doc_id = metadata.get("document_id")
                    if doc_id and doc_id not in documents:
                        documents[doc_id] = {
                            "document_id": doc_id,
                            "title": metadata.get("title", "Unknown"),
                            "source_type": metadata.get("source_type", "unknown"),
                            "source_url": metadata.get("source_url", ""),
                            "ingested_at": metadata.get("ingested_at", ""),
                        }
            
            return list(documents.values())
            
        except Exception as e:
            logger.error(f"Error listing documents: {e}")
            return []
    
    def get_stats(self) -> dict:
        """Get statistics about the vector store."""
        try:
            collection = self.vector_store._collection
            count = collection.count()
            documents = self.list_documents()
            
            return {
                "total_chunks": count,
                "total_documents": len(documents),
                "persist_directory": self.config.persist_directory,
                "collection_name": self.config.collection_name,
            }
            
        except Exception as e:
            logger.error(f"Error getting stats: {e}")
            return {"error": str(e)}

    # =========================================================================
    # CONVENIENCE METHODS FOR TOKEN SERVER API
    # =========================================================================
    
    def index_youtube(self, url: str, title: Optional[str] = None) -> dict:
        """
        Ingest a YouTube video by URL.
        
        Args:
            url: YouTube video URL
            title: Optional custom title
            
        Returns:
            Dictionary with ingestion results
        """
        result = self.ingest(url, title=title)
        if result.get("success"):
            result["chunk_count"] = result.get("num_chunks", 0)
        return result
    
    def index_web(self, url: str, title: Optional[str] = None) -> dict:
        """
        Ingest a web article by URL.
        
        Args:
            url: Web page URL
            title: Optional custom title
            
        Returns:
            Dictionary with ingestion results
        """
        result = self.ingest(url, title=title)
        if result.get("success"):
            result["chunk_count"] = result.get("num_chunks", 0)
        return result
    
    def index_pdf(self, source: str, title: Optional[str] = None) -> dict:
        """
        Ingest a PDF from URL or file path.
        
        Args:
            source: PDF URL or file path
            title: Optional custom title
            
        Returns:
            Dictionary with ingestion results
        """
        result = self.ingest(source, title=title)
        if result.get("success"):
            result["chunk_count"] = result.get("num_chunks", 0)
        return result
    
    def index_pdf_bytes(self, content: bytes, filename: str) -> dict:
        """
        Ingest a PDF from raw bytes (file upload).
        
        Args:
            content: PDF file bytes
            filename: Original filename
            
        Returns:
            Dictionary with ingestion results
        """
        import tempfile
        import os
        
        # Write to temp file and ingest
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        
        try:
            result = self.ingest(tmp_path, title=filename.replace(".pdf", "").replace("_", " "))
            if result.get("success"):
                result["chunk_count"] = result.get("num_chunks", 0)
            return result
        finally:
            # Clean up temp file
            if os.path.exists(tmp_path):
                os.remove(tmp_path)
    
    def index_docx_bytes(self, content: bytes, filename: str) -> dict:
        """
        Ingest a DOCX file from raw bytes.
        
        Args:
            content: DOCX file bytes
            filename: Original filename
            
        Returns:
            Dictionary with ingestion results
        """
        try:
            from docx import Document as DocxDocument
            import io
            
            # Parse DOCX
            doc = DocxDocument(io.BytesIO(content))
            
            # Extract text from paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            full_text = "\n\n".join(text_parts)
            title = filename.replace(".docx", "").replace("_", " ")
            
            result = self.ingest_text(full_text, title)
            if result.get("success"):
                result["chunk_count"] = result.get("num_chunks", 0)
            return result
            
        except ImportError:
            logger.error("python-docx not installed")
            return {"success": False, "error": "DOCX support not available"}
        except Exception as e:
            logger.error(f"Error indexing DOCX: {e}")
            return {"success": False, "error": str(e)}
    
    def index_text(self, text: str, title: str) -> dict:
        """
        Ingest raw text content.
        
        Args:
            text: Text content
            title: Document title
            
        Returns:
            Dictionary with ingestion results
        """
        result = self.ingest_text(text, title)
        if result.get("success"):
            result["chunk_count"] = result.get("num_chunks", 0)
        return result
